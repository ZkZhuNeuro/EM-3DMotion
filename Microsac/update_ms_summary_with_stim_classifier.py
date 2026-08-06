"""Add the vertical-MS Stim/NonStim classifier to the concise DOCX summary."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table
from docx.shared import Inches


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "MS_choice_stimulation_summary.docx"
OUTPUT = ROOT / "MS_choice_stimulation_summary_updated.docx"
FIGURE = Path(
    r"C:\EM\Microsac\population_merged_12ms_no_smoothing"
    r"\population_analysis\stim_from_vertical_ms"
    r"\stim_from_vertical_ms_summary.png"
)


def paragraph_with_text(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text == exact_text:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact_text}")


def replace_run_text(run_element, text):
    """Replace a run's displayed text while retaining its run properties."""
    for child in list(run_element):
        if child.tag != qn("w:rPr"):
            run_element.remove(child)
    text_element = run_element.makeelement(qn("w:t"))
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)


def clone_single_run_paragraph(template, text):
    element = deepcopy(template._p)
    runs = element.findall(qn("w:r"))
    if not runs:
        raise ValueError(f"Template paragraph has no direct run: {template.text}")
    replace_run_text(runs[0], text)
    for extra in runs[1:]:
        element.remove(extra)
    return element


def clone_labeled_paragraph(template, label, body):
    element = deepcopy(template._p)
    runs = element.findall(qn("w:r"))
    if len(runs) < 2:
        raise ValueError(f"Template paragraph needs two runs: {template.text}")
    replace_run_text(runs[0], label)
    replace_run_text(runs[1], body)
    for extra in runs[2:]:
        element.remove(extra)
    return element


def replace_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = text
        for extra in paragraph.runs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        paragraph.add_run(text)


def main():
    if not SOURCE.is_file():
        raise FileNotFoundError(SOURCE)
    if not FIGURE.is_file():
        raise FileNotFoundError(FIGURE)

    document = Document(SOURCE)

    # Update the document timestamp.
    prepared = paragraph_with_text(
        document,
        "Prepared from repository analyses and saved outputs | 21 July 2026",
    )
    prepared.runs[0].text = (
        "Prepared from repository analyses and saved outputs | 28 July 2026"
    )

    # Extend the lead conclusion without changing the existing callout design.
    bottom_line = next(
        p for p in document.paragraphs if p.text.startswith("BOTTOM LINE")
    )
    bottom_line.runs[1].text = (
        "Across 77,742 trials containing at least one microsaccade, "
        "electrical stimulation neither measurably altered how trial-mean "
        "MS displacement predicted choice nor left a useful vertical-MS "
        "signature of Stim versus NonStim trials. The only detectable "
        "reverse-model association (Clay-FST) had near-chance held-out "
        "discrimination."
    )

    # Templates retain the document's established styles and direct formatting.
    heading_template = paragraph_with_text(
        document, "Session-level checks and interpretation"
    )
    intro_template = paragraph_with_text(
        document,
        "The pooled models are the clearest population-level answer: they use "
        "all usable trials while preserving within-session MS scaling and "
        "separate NonStim/Stim baselines for each recording.",
    )
    caption_template = next(
        p for p in document.paragraphs if p.text.startswith("Figure 1.")
    )
    session_caption = next(
        p for p in document.paragraphs if p.text.startswith("Figure 2.")
    )
    table_title_template = next(
        p for p in document.paragraphs if p.text.startswith("Table 1.")
    )
    note_template = next(
        p for p in document.paragraphs if p.text.startswith("Note: Interaction")
    )
    result_template = next(
        p for p in document.paragraphs if p.text.startswith("Result.")
    )

    target = heading_template
    target._p.addprevious(
        clone_single_run_paragraph(
            heading_template, "Can vertical MS identify a Stim trial?"
        )
    )
    target._p.addprevious(
        clone_single_run_paragraph(
            intro_template,
            "Reverse model: trial-average vertical MS was standardized within "
            "session. We fit Stim ~ SessionID + MeanMSY_Z for association and "
            "Stim ~ MeanMSY_Z for prediction, evaluated by leave-one-session-"
            "out AUC. Holm correction covered the four monkey-area groups.",
        )
    )

    figure_paragraph = document.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_run = figure_paragraph.add_run()
    shape = figure_run.add_picture(str(FIGURE), width=Inches(6.00))
    shape._inline.docPr.set(
        "descr",
        "Held-out session AUCs and session-adjusted logistic coefficients for "
        "predicting Stim versus NonStim from vertical microsaccade displacement.",
    )
    target._p.addprevious(figure_paragraph._p)

    target._p.addprevious(
        clone_single_run_paragraph(
            caption_template,
            "Figure 2. Trial-level discrimination of Stim versus NonStim from "
            "vertical MS. Left: mean held-out-session AUC. Right: "
            "session-adjusted log-odds slope per within-session SD. Bars are "
            "95% confidence intervals.",
        )
    )
    target._p.addprevious(
        clone_single_run_paragraph(
            table_title_template,
            "Table 2. Stim classification from vertical MS",
        )
    )

    table_values = [
        [
            "Group",
            "Sessions / trials",
            "Slope (Holm p)",
            "Held-out AUC (95% CI)",
            "Balanced acc.",
        ],
        ["Jim-MT", "56 / 13,346", "-0.011 (1.000)", ".507 (.487-.528)", ".488"],
        ["Jim-FST", "94 / 18,144", "-0.002 (1.000)", ".474 (.460-.488)", ".500"],
        ["Clay-MT", "38 / 16,462", "-0.011 (1.000)", ".506 (.486-.527)", ".499"],
        ["Clay-FST", "65 / 29,790", "-0.040 (.0026)", ".514 (.495-.533)", ".505"],
    ]
    table_element = deepcopy(document.tables[0]._tbl)
    classifier_table = Table(table_element, document._body)
    for row, values in zip(classifier_table.rows, table_values):
        for cell, value in zip(row.cells, values):
            replace_cell_text(cell, value)
    target._p.addprevious(table_element)

    target._p.addprevious(
        clone_single_run_paragraph(
            note_template,
            "Note: Slope is the session-adjusted log-odds coefficient per "
            "within-session SD; p values are Holm-adjusted. AUC = 0.5 is "
            "chance.",
        )
    )
    target._p.addprevious(
        clone_labeled_paragraph(
            result_template,
            "Result. ",
            "Vertical MS did not usefully distinguish individual Stim from "
            "NonStim trials. Clay-FST had a detectable but tiny association "
            "(slope = -0.0396; odds ratio = 0.961; raw difference = -0.0046 "
            "degrees), yet held-out AUC was 0.514 and balanced accuracy was "
            "0.505. Jim-FST's below-chance AUC indicates an unstable direction "
            "across sessions, not a useful classifier.",
        )
    )

    # Integrate the new result into the closing interpretation and sources.
    inference = next(
        p for p in document.paragraphs if p.text.startswith("Inference.")
    )
    inference.runs[1].text += (
        " Vertical MS also did not reliably identify whether stimulation was "
        "present on an individual trial."
    )
    session_caption.runs[0].text = session_caption.runs[0].text.replace(
        "Figure 2.", "Figure 3.", 1
    )
    sources = next(
        p for p in document.paragraphs if p.text.startswith("Analysis sources:")
    )
    sources.runs[0].text = (
        "Analysis sources: README.md; analyze_microsaccades.m; "
        "analyze_trialwise_ms_choice.m; analyze_grouped_ms_choice.m; "
        "analyze_paired_ms_bias_lme.m; analyze_stim_from_vertical_ms.m; and "
        "saved session/group/classifier summary CSV files under "
        r"C:\EM\Microsac\population_merged_12ms_no_smoothing"
        r"\population_analysis."
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
