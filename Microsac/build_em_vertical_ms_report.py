from __future__ import annotations

import math
import os
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(r"C:\Users\zzhu329\Documents\GitHub\EM-3DMotion\Microsac")
POP_ROOT = Path(r"C:\EM\Microsac\population_merged_12ms_no_smoothing")
ANALYSIS_ROOT = POP_ROOT / "population_analysis"
OUTPUT_DOCX = REPO_ROOT / "EM_vertical_MS_choice_accountability_report.docx"

SKILL_ROOT = Path(
    r"C:\Users\zzhu329\.codex\plugins\cache\openai-primary-runtime"
    r"\documents\26.723.12215\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


# Resolved design system:
#   Preset: standard_business_brief
#   Header pattern: editorial_cover
# Named overrides:
#   - report_title: 29 pt navy, centered (cover-only)
#   - scientific_table: 8.5 pt body / 8.5 pt bold header
#   - figure_caption: 9 pt gray italic, 1.05 line spacing
#   - result_callout: 1-column, pale blue fill with navy left rule
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "5B6573"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "EAF2F8"
LIGHT_GOLD = "FFF7E6"
GREEN = "2F6F4E"
RED = "9B1C1C"
BLACK = "000000"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge in edges.items():
        tag = f"w:{edge_name}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge:
                element.set(qn(f"w:{key}"), str(edge[key]))


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color=BLACK, size=8.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    alignments: list[int | None] | None = None,
    header_fill: str = LIGHT_GRAY,
    font_size: float = 8.5,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].height = None
    mark_header_row(table.rows[0])
    if alignments is None:
        alignments = [None] * len(headers)
    for i, text in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            text,
            bold=True,
            color=NAVY,
            size=font_size,
            align=alignments[i],
        )
        set_cell_shading(table.rows[0].cells[i], header_fill)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(values):
            set_cell_text(
                row.cells[i],
                value,
                size=font_size,
                align=alignments[i],
            )
            if row_idx % 2 == 1:
                set_cell_shading(row.cells[i], "FAFBFC")
    apply_table_geometry(
        table,
        widths_dxa,
        table_width_dxa=sum(widths_dxa),
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS,
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(3)
    return table


def add_source_note(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Source Note"]
    p.add_run(text)


def add_callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = LIGHT_BLUE,
    border_color: str = BLUE,
):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.04)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), border_color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)
    return p


def add_caption(doc: Document, number: int, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Figure Caption"]
    r = p.add_run(f"Figure {number}. ")
    set_run_font(r, size=9, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=9, color=GRAY, italic=True)
    return p


def add_figure(
    doc: Document,
    path: Path,
    number: int,
    caption: str,
    alt_text: str,
    *,
    width: float = 6.35,
    page_break_before: bool = False,
):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", f"Figure {number}")
    add_caption(doc, number, caption)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, size=10.5, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_lead) :])
        set_run_font(r, size=10.5, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=BLACK)
    return p


def add_code_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles["Code"]
    p.add_run(text)
    return p


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=GRAY)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in [s.name for s in styles]:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(GRAY)
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.05

    if "Source Note" not in [s.name for s in styles]:
        source = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = styles["Source Note"]
    source.font.name = "Calibri"
    source._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    source._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    source.font.size = Pt(8.5)
    source.font.italic = True
    source.font.color.rgb = rgb(GRAY)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(6)
    source.paragraph_format.line_spacing = 1.05

    if "Code" not in [s.name for s in styles]:
        code = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.5)
    code.font.color.rgb = rgb(DARK_BLUE)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(1)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0

    # Quiet running header and page number; first page remains uncluttered.
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header_p.add_run("TECHNICAL ANALYSIS REPORT  |  VERTICAL MS AND CHOICE")
    set_run_font(r, size=8.2, color=GRAY, bold=True)
    header_p.paragraph_format.space_after = Pt(0)
    first_header = section.first_page_header.paragraphs[0]
    first_header.text = ""

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    add_page_number(footer_p)
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.text = ""


def fmt(value, digits=3):
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "NA"
    return f"{float(value):.{digits}f}"


def fmt_p(value):
    value = float(value)
    if math.isnan(value):
        return "NA"
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.3f}"


def fmt_ci(low, high, digits=3):
    return f"[{float(low):.{digits}f}, {float(high):.{digits}f}]"


def load_results():
    hist = pd.read_csv(
        ANALYSIS_ROOT / "histograms" / "population_ms_metric_histogram_summary.csv"
    )
    stim_tests = pd.read_csv(
        POP_ROOT / "population_microsaccade_stim_nonstim_tests_complete.csv"
    )
    paired = pd.read_csv(
        ANALYSIS_ROOT
        / "paired_ms_y_bias_lme"
        / "paired_ms_y_bias_lme_summary.csv"
    )
    centered = pd.read_csv(
        ANALYSIS_ROOT
        / "paired_ms_y_bias_lme"
        / "centered_change"
        / "centered_delta_ms_y_bias_test_summary.csv"
    )
    sessions = pd.read_csv(
        ANALYSIS_ROOT
        / "trialwise_ms_choice"
        / "trialwise_ms_choice_session_summary.csv"
    )
    grouped = pd.read_csv(
        ANALYSIS_ROOT
        / "grouped_monkey_roi_ms_choice"
        / "grouped_ms_choice_summary.csv"
    )
    pooled = pd.read_csv(
        ANALYSIS_ROOT
        / "grouped_monkey_roi_ms_choice"
        / "pooled_logistic"
        / "pooled_logistic_ms_choice_summary.csv"
    )
    comparison = pd.read_csv(
        ANALYSIS_ROOT
        / "grouped_monkey_roi_ms_choice"
        / "pooled_logistic"
        / "pooled_vs_session_stratified_summary.csv"
    )
    return hist, stim_tests, paired, centered, sessions, grouped, pooled, comparison


def build_report():
    hist, stim_tests, paired, centered, sessions, grouped, pooled, comparison = (
        load_results()
    )
    doc = Document()
    configure_document(doc)

    # Cover.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(74)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    r = kicker.add_run("TECHNICAL ANALYSIS REPORT")
    set_run_font(r, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("Can EM-Driven Vertical Microsaccades\nAccount for Up/Down Choices?")
    set_run_font(r, size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run(
        "Multilevel evidence from microsaccade generation, session-level bias, "
        "trial-level choice, and MS-by-EM interaction models"
    )
    set_run_font(r, size=13.5, color=GRAY, italic=True)

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(22)
    r = metadata.add_run(
        "Jim and Clay  |  MT and FST  |  12 ms detector  |  24 July 2026"
    )
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)

    add_callout(
        doc,
        "Bottom line",
        "EM-related changes in mean vertical MS displacement covary with "
        "EM-related changes in fitted behavioral bias across sessions, most "
        "convincingly in MT. However, vertical MS does not robustly predict the "
        "monkey's individual up/down choice after session and visual-evidence "
        "control, and EM does not reliably change that slope. The current data "
        "therefore support vertical MS as a session-level correlate or marker, "
        "not as a demonstrated mediator of EM's behavioral effect.",
    )

    scope = doc.add_paragraph()
    scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scope.paragraph_format.space_before = Pt(9)
    r = scope.add_run(
        "Primary trial-level sample: 253 sessions and 77,742 trials containing at least one MS"
    )
    set_run_font(r, size=9.5, color=GRAY)
    doc.add_page_break()

    # Executive conclusion.
    add_heading(doc, "1. Executive conclusion", 1)
    add_body(
        doc,
        "The biological question is a mediation question: did EM alter vertical "
        "microsaccades, and did those EM-driven MS changes in turn produce the "
        "up/down choice bias? Because the choice targets are vertical, MS_y - "
        "the signed vertical displacement, positive upward - is the primary "
        "oculomotor variable. Three links are needed for a persuasive account.",
    )
    evidence_rows = [
        [
            "A. EM -> MS_y",
            "EM must reproducibly move the vertical-MS distribution.",
            "Group means changed only slightly and in mixed directions; no common vertical shift.",
            "Weak / heterogeneous",
        ],
        [
            "B. DeltaMS_y -> DeltaBias",
            "Sessions with upward/downward MS changes should show matching behavioral-bias changes.",
            "Sign concordance: MT 69.2%, FST 65.2%; FDR-significant. Magnitude slope positive only in MT.",
            "Supportive at session level",
        ],
        [
            "C. MS_y -> trial choice",
            "Within a session, vertical MS direction should predict the later binary choice.",
            "No adjusted group slope or joint MS effect survived correction; no session fit survived Holm.",
            "Not supported",
        ],
        [
            "D. EM x MS_y",
            "If EM changes MS-choice coupling, the interaction should differ from zero.",
            "All vertical and 2-D interactions were nonsignificant after correction.",
            "Not supported",
        ],
        [
            "Overall mediation",
            "A formal indirect EM -> MS_y -> choice effect should be estimable and nonzero.",
            "No formal multilevel mediation was fit; the necessary trial-level b-path is not established.",
            "Not established",
        ],
    ]
    add_table(
        doc,
        ["Claim", "Required evidence", "Observed evidence", "Verdict"],
        evidence_rows,
        [1550, 2450, 3650, 1710],
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        font_size=8.2,
    )
    add_callout(
        doc,
        "Important distinction",
        "The absence of an MS_y-by-EM interaction does not by itself rule out "
        "mediation; mediation can occur with a common MS_y slope in both "
        "conditions. Here, the more consequential result is that the adjusted "
        "within-session MS_y-to-choice slope itself is not reliably different "
        "from zero.",
        fill=LIGHT_GOLD,
        border_color="C58B00",
    )

    # Data and preprocessing.
    add_heading(doc, "2. Data and preprocessing", 1)
    add_heading(doc, "2.1 Behavioral and oculomotor definitions", 2)
    add_body(
        doc,
        "Choice was the decoded binary response (0/1) for the monkey's two "
        "vertical choice targets. The models estimate P(Choice = 1); the sign "
        "of a coefficient should be translated to 'up' or 'down' only after "
        "confirming the task's response-code mapping. MS_y is signed vertical "
        "eye displacement in degrees, with positive values defined as upward.",
    )
    add_body(
        doc,
        "Microsaccades were detected only during the visual-stimulus interval "
        "from event ID 118 to event ID 130, before the instructed choice "
        "saccade. The completed population batch used a 12 ms minimum duration, "
        "no smoothing, and the merged/unmatched-eye detector. For a trial with "
        "multiple MS events, event displacement vectors were averaged within "
        "that trial.",
    )
    doc.add_page_break()
    add_heading(doc, "2.2 Analysis samples", 2)
    group_order = ["Jim-MT", "Jim-FST", "Clay-MT", "Clay-FST"]
    grouped = grouped.set_index("Group").loc[group_order].reset_index()
    sample_rows = []
    for _, row in grouped.iterrows():
        sample_rows.append(
            [
                row["Group"],
                f"{int(row['NSessions'])}",
                f"{int(row['NTrials']):,}",
                f"{int(row['NNonStimTrials']):,}",
                f"{int(row['NStimTrials']):,}",
            ]
        )
    sample_rows.append(
        [
            "Total",
            "253",
            f"{int(grouped.NTrials.sum()):,}",
            f"{int(grouped.NNonStimTrials.sum()):,}",
            f"{int(grouped.NStimTrials.sum()):,}",
        ]
    )
    add_table(
        doc,
        ["Monkey-area", "Sessions", "MS trials", "NonStim", "Stim"],
        sample_rows,
        [2100, 1300, 1950, 1950, 2060],
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
        ],
    )
    add_source_note(
        doc,
        "Source: grouped_ms_choice_summary.csv. The trialwise session analysis "
        "attempted 254 sessions; 246 met all per-session fitting criteria. The "
        "grouped table contains 253 sessions with usable pooled trial rows.",
    )
    add_body(
        doc,
        "Trial-level models were conditional on observing at least one MS. This "
        "is appropriate for asking whether the direction of an observed MS "
        "predicts choice, but it does not estimate behavioral pathways through "
        "EM-induced changes in the probability of having an MS.",
    )

    # Models.
    add_heading(doc, "3. Statistical models", 1)
    add_body(
        doc,
        "The repository contains complementary models at three levels. They "
        "answer different questions and should not be treated as interchangeable.",
    )
    model_rows = [
        [
            "Stim/NonStim MS comparison",
            "Per session",
            "Scalar differences and 2-D first-moment distance; 1,000 label permutations within visual-condition x signed-coherence strata",
            "Does EM change MS production or kinematics in that session?",
        ],
        [
            "Paired bias LME",
            "Session-condition",
            "Bias ~ MS_y + condition + (1|SessionKey); interaction sensitivity adds MS_y x condition",
            "Do session means covary after pairing Stim/NonStim conditions?",
        ],
        [
            "Centered change",
            "Session",
            "DeltaMS_y and DeltaBias; exact one-sided quadrant-concordance test; positive origin-slope sensitivity",
            "Do EM-induced MS_y and bias changes have matching signs/magnitudes?",
        ],
        [
            "Session logistic",
            "Trial within session",
            "Simple, visual-evidence-adjusted vertical, and adjusted X/Y models; Jeffreys-prior binomial fits",
            "Does MS direction predict choice in an individual recording session?",
        ],
        [
            "Grouped stratified logistic",
            "Trial pooled within monkey-area",
            "Session-specific Stim/NonStim intercepts plus MS_y x Stim; adjusted model adds signed coherence x visual condition",
            "Is there a common within-session MS_y effect or EM interaction?",
        ],
        [
            "Exact pooled logistic",
            "All trial rows in monkey-area",
            "Choice ~ MeanMSY_Z * StimCondition; no SessionID term",
            "Sensitivity to the requested unstratified pooled specification",
        ],
    ]
    add_table(
        doc,
        ["Model", "Unit", "Specification", "Question"],
        model_rows,
        [1800, 1200, 3550, 2810],
        alignments=[None, None, None, None],
        font_size=7.9,
    )

    add_heading(doc, "3.1 Session-level EM effects on MS pattern", 2)
    add_body(
        doc,
        "For each session, stimulation labels were shuffled 1,000 times within "
        "visual-condition and signed-coherence strata. Scalar outcomes were MS "
        "rate, probability of any MS, per-trial mean amplitude, peak velocity, "
        "and duration. Direction was tested as the Euclidean distance between "
        "Stim and NonStim two-dimensional mean displacement vectors. Event "
        "properties were summarized within trials before testing so that trials "
        "with many events did not count as independent observations.",
    )

    add_heading(doc, "3.2 Session-mean behavioral-bias models", 2)
    add_code_line(doc, "Bias ~ MS_y + condition + (1|SessionKey)")
    add_code_line(doc, "Bias ~ MS_y * condition + (1|SessionKey)")
    add_body(
        doc,
        "Only complete Stim/NonStim session pairs with acceptable psychometric "
        "fits were used. The common-slope model asks whether mean vertical MS "
        "displacement covaries with fitted behavioral bias across the paired "
        "conditions. The interaction model asks whether that slope differs "
        "under EM. Benjamini-Hochberg FDR correction was applied across the four "
        "cue definitions separately within each ROI.",
    )

    add_heading(doc, "3.3 Centered EM-induced change analysis", 2)
    add_code_line(doc, "DeltaMS_y = MS_y(Stim) - MS_y(NonStim)")
    add_code_line(doc, "DeltaBias = Bias(Stim) - Bias(NonStim)")
    add_body(
        doc,
        "The primary exact binomial test asks whether DeltaMS_y and DeltaBias "
        "have the same sign more often than 50%. A secondary one-sided regression "
        "fits DeltaBias = beta * DeltaMS_y through the origin. FDR correction was "
        "reported within ROI and across all eight ROI-by-cue tests. The sign test "
        "is robust to scale, whereas the origin regression tests whether larger "
        "MS changes correspond to proportionally larger bias changes.",
    )

    add_heading(doc, "3.4 Trial-level logistic models", 2)
    add_code_line(doc, "Simple:   Choice ~ MeanMSY_Z * StimCondition")
    add_code_line(
        doc,
        "Adjusted: Choice ~ SignedCoherence * VisualCondition + StimCondition + MeanMSY_Z * StimCondition",
    )
    add_code_line(
        doc,
        "Vector:   Adjusted base + MeanMSX_Z * StimCondition + MeanMSY_Z * StimCondition",
    )
    add_body(
        doc,
        "MeanMSY_Z and MeanMSX_Z were standardized within recording session; a "
        "slope is therefore the change in log odds of Choice = 1 per one "
        "within-session SD of trial-average MS displacement. Jeffreys-prior "
        "bias reduction was used to avoid infinite estimates in sparse "
        "choice-by-cue cells. Session fits required at least 30 usable MS trials, "
        "at least 10 in each stimulation condition, variation in MS_y, and both "
        "choices within each condition.",
    )
    add_body(
        doc,
        "For the grouped analysis, trials were pooled into Jim-MT, Jim-FST, "
        "Clay-MT, and Clay-FST, but SessionID-by-StimCondition terms supplied "
        "session-specific NonStim and Stim intercepts. The primary adjusted "
        "group formula was:",
    )
    add_code_line(
        doc,
        "Choice ~ SessionID * StimCondition + SignedCoherence * VisualCondition + MeanMSY_Z * StimCondition",
    )
    add_body(
        doc,
        "The primary MS test was the average of the NonStim and Stim vertical "
        "slopes; the primary moderation test was the Stim-minus-NonStim slope "
        "difference. Joint tests evaluated any vertical-MS effect and, in the "
        "vector model, the combined horizontal/vertical effects. Holm-Bonferroni "
        "correction was applied across the four monkey-area groups separately "
        "for each model/test family.",
    )
    add_heading(doc, "3.5 Pooled sensitivity model", 2)
    add_code_line(doc, "Choice ~ MeanMSY_Z * StimCondition")
    add_body(
        doc,
        "This exact pooled model used the same trial rows and within-session "
        "standardization but omitted SessionID. It treats trial rows as "
        "independent and can attribute stable session-to-session differences in "
        "choice baseline to MS_y. It is retained as a sensitivity analysis; the "
        "session-stratified grouped model is the preferred inferential model.",
    )

    # Results: EM on MS_y.
    add_heading(doc, "4. Results", 1)
    add_heading(doc, "4.1 Did EM change vertical MS?", 2)
    vertical = hist[hist.Metric == "MS_y"].copy()
    vertical["Group"] = vertical.Subpopulation.str.replace("_", "-", regex=False)
    vertical = vertical.set_index(["Group", "Condition"])
    vertical_rows = []
    direction_counts = (
        stim_tests[stim_tests.Metric == "DirectionFirstMoment"]
        .assign(Group=lambda d: d.Monkey.astype(str) + "-" + d.ROI.astype(str))
        .groupby("Group")
        .agg(
            NSessions=("PermutationP", "count"),
            RawSig=("PermutationP", lambda x: int((x < 0.05).sum())),
        )
    )
    for group in group_order:
        ns = vertical.loc[(group, "NonStim")]
        st = vertical.loc[(group, "Stim")]
        diff = st["Mean"] - ns["Mean"]
        dc = direction_counts.loc[group]
        vertical_rows.append(
            [
                group,
                f"{int(ns['N'])}",
                fmt(ns["Mean"], 5),
                fmt(st["Mean"], 5),
                f"{diff:+.5f}",
                f"{ns['NegativePercent']:.1f}% / {st['NegativePercent']:.1f}%",
                f"{int(dc.RawSig)}/{int(dc.NSessions)}",
            ]
        )
    add_table(
        doc,
        [
            "Group",
            "N",
            "NonStim MS_y",
            "Stim MS_y",
            "Mean delta",
            "Negative % N/S",
            "2-D direction p<.05",
        ],
        vertical_rows,
        [1200, 650, 1370, 1370, 1200, 1650, 1920],
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        font_size=8.0,
    )
    add_source_note(
        doc,
        "MS_y values are session-condition means in degrees. The last column is "
        "the number of individual sessions with an unadjusted p < .05 in the "
        "stratified 2-D mean-vector permutation test; it is not a vertical-only "
        "test and should not be read as a group-level effect.",
    )
    add_body(
        doc,
        "The mean Stim-minus-NonStim vertical change was small and did not have "
        "a common sign across groups: -0.00047 degrees in Jim-MT, +0.00036 in "
        "Jim-FST, +0.00187 in Clay-MT, and -0.00845 in Clay-FST. Clay's MS_y "
        "distribution was strongly downward in both conditions, whereas Jim's "
        "was much closer to zero. Thus EM did not impose a single population-wide "
        "vertical shift. Individual sessions did show 2-D pattern changes, "
        "especially in Jim, but a change in the two-dimensional vector need not "
        "be a behaviorally aligned change in MS_y.",
    )
    add_figure(
        doc,
        ANALYSIS_ROOT
        / "histograms"
        / "ms_vertical_displacement_histograms_by_subpopulation.png",
        1,
        "Session-level distributions of mean vertical MS displacement by "
        "monkey-area group and EM condition. Solid lines mark condition medians; "
        "the dotted line marks zero. The extensive overlap and mixed direction "
        "of mean shifts argue against a uniform EM-induced vertical-MS change.",
        "Four histograms comparing NonStim and Stim mean vertical microsaccade "
        "displacement in Jim-MT, Jim-FST, Clay-MT, and Clay-FST.",
        width=6.35,
    )

    # Session-level bias/change results.
    add_heading(doc, "4.2 Do EM-induced MS_y changes track behavioral-bias changes?", 2)
    combined_centered = centered[centered.CueName == "Combined"].copy()
    change_rows = []
    for roi_name in ["MT", "FST"]:
        row = combined_centered[combined_centered.ROI == roi_name].iloc[0]
        change_rows.append(
            [
                roi_name,
                f"{int(row.NSessions)}",
                f"{int(row.ConcordantCount)}/{int(row.NForQuadrantTest)} "
                f"({100 * row.ConcordanceRate:.1f}%)",
                fmt_ci(row.ConcordanceCI95Lower, row.ConcordanceCI95Upper),
                fmt_p(row.QuadrantP_Q_FDR_All),
                f"{row.OriginSlope:.3f} +/- {row.OriginSlopeSE:.3f}",
                fmt_p(row.OriginSlopeP_Q_FDR_All),
            ]
        )
    add_table(
        doc,
        [
            "ROI",
            "Sessions",
            "Concordant",
            "95% CI",
            "Quadrant q",
            "Origin slope +/- SE",
            "Slope q",
        ],
        change_rows,
        [850, 950, 1750, 1500, 1150, 1950, 1210],
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        font_size=8.1,
    )
    add_body(
        doc,
        "For the combined-cue behavioral bias, MT showed matching signs of "
        "DeltaMS_y and DeltaBias in 63 of 91 sessions (69.2%, FDR q = "
        "0.00042). The positive through-origin slope was also significant "
        "(5.09 +/- 1.79, q = 0.0061), so sessions with larger vertical-MS shifts "
        "tended to have larger bias shifts in the same direction.",
    )
    add_body(
        doc,
        "FST also showed significant sign concordance: 103 of 158 non-axis "
        "sessions (65.2%, q = 0.00038). However, its origin slope was small and "
        "nonsignificant (0.27 +/- 0.61, q = 0.328). The FST evidence therefore "
        "supports directional agreement but not proportional scaling of effect "
        "magnitude. Across cue-specific analyses, quadrant concordance remained "
        "FDR-significant in all MT and FST cue definitions; origin slopes were "
        "significant in all MT definitions and in none of the FST definitions.",
    )
    add_figure(
        doc,
        ANALYSIS_ROOT
        / "paired_ms_y_bias_lme"
        / "centered_change"
        / "mt_centered_delta_ms_y_bias.png",
        2,
        "MT session-level EM-induced changes. Green points fall in concordant "
        "quadrants, where DeltaMS_y and DeltaBias share a sign. The combined and "
        "cue-specific results show both excess sign concordance and positive "
        "magnitude scaling.",
        "Four MT scatterplots of stimulation-induced change in vertical "
        "microsaccade displacement versus change in behavioral bias.",
        width=6.35,
        page_break_before=True,
    )
    add_figure(
        doc,
        ANALYSIS_ROOT
        / "paired_ms_y_bias_lme"
        / "centered_change"
        / "fst_centered_delta_ms_y_bias.png",
        3,
        "FST session-level EM-induced changes. Sign concordance exceeds 50%, "
        "but the fitted through-origin slopes are shallow and nonsignificant, "
        "indicating weak correspondence in change magnitude.",
        "Four FST scatterplots of stimulation-induced change in vertical "
        "microsaccade displacement versus change in behavioral bias.",
        width=6.35,
        page_break_before=True,
    )

    add_heading(doc, "4.3 Session-mean mixed models", 2)
    paired_combined = paired[paired.CueName == "Combined"]
    lme_rows = []
    for roi_name in ["MT", "FST"]:
        row = paired_combined[paired_combined.ROI == roi_name].iloc[0]
        lme_rows.append(
            [
                roi_name,
                f"{int(row.NSessions)}",
                f"{row.CommonSlope:.3f} +/- {row.CommonSlopeSE:.3f}",
                fmt_p(row.CommonSlopeQ_FDR),
                f"{row.InteractionSlopeDifference:.3f}",
                fmt_p(row.InteractionQ_FDR),
            ]
        )
    add_table(
        doc,
        [
            "ROI",
            "Sessions",
            "Common MS_y slope +/- SE",
            "Slope q",
            "MS_y x EM",
            "Interaction q",
        ],
        lme_rows,
        [900, 1100, 2500, 1150, 1650, 2060],
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    add_body(
        doc,
        "The paired session-mean LME gave a positive common MS_y-bias slope for "
        "the combined cue in MT (1.354 +/- 0.410, q = 0.0047) and FST "
        "(0.744 +/- 0.197, q = 0.00075). The MS_y-by-condition interaction was "
        "not significant in either ROI (MT q = 0.901; FST q = 0.762). These "
        "results agree with a stable session-level association but do not show "
        "that the MS occurring on a particular trial determines the subsequent "
        "choice.",
    )

    # Trial-level grouped results.
    doc.add_page_break()
    add_heading(doc, "4.4 Does vertical MS predict the individual up/down choice?", 2)
    grouped_rows = []
    for _, row in grouped.iterrows():
        avg_low = row.AdjustedAverageSlope - 1.96 * row.AdjustedAverageSlopeSE
        avg_high = row.AdjustedAverageSlope + 1.96 * row.AdjustedAverageSlopeSE
        int_low = row.AdjustedInteraction - 1.96 * row.AdjustedInteractionSE
        int_high = row.AdjustedInteraction + 1.96 * row.AdjustedInteractionSE
        grouped_rows.append(
            [
                row.Group,
                f"{row.AdjustedAverageSlope:.3f} "
                f"[{avg_low:.3f}, {avg_high:.3f}]",
                fmt_p(row.AdjustedAverageSlopeP_Holm),
                f"{row.AdjustedNonStimSlope:.3f} / {row.AdjustedStimSlope:.3f}",
                f"{row.AdjustedInteraction:.3f} "
                f"[{int_low:.3f}, {int_high:.3f}]",
                fmt_p(row.AdjustedInteractionP_Holm),
                fmt_p(row.VectorInteractionP_Holm),
            ]
        )
    add_table(
        doc,
        [
            "Group",
            "Adjusted average slope [95% CI]",
            "Avg Holm p",
            "N/S slopes",
            "MS_y x EM [95% CI]",
            "Int Holm p",
            "2-D int Holm p",
        ],
        grouped_rows,
        [1150, 2050, 1000, 1400, 1900, 950, 910],
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        font_size=7.8,
    )
    add_body(
        doc,
        "No adjusted average vertical-MS slope was significant after correction "
        "(Holm p = 0.604 to 1.000). The largest condition-specific coefficient "
        "was Clay-MT in NonStim trials (0.051 log-odds per within-session SD, "
        "unadjusted p = 0.078), but its Stim slope was near zero (0.005, "
        "p = 0.839), and the interaction was not significant. Effect sizes were "
        "small: the adjusted average odds ratios per one-SD MS_y increase were "
        "approximately 1.008 to 1.029.",
    )
    complete = sessions[sessions.Status == "Complete"].copy()
    n_complete = len(complete)
    adjusted_any_raw = int((complete.VerticalAnyEffectJointP < 0.05).sum())
    adjusted_int_raw = int((complete.VerticalInteractionP < 0.05).sum())
    vector_any_raw = int((complete.VectorAnyEffectJointP < 0.05).sum())
    vector_int_raw = int((complete.VectorInteractionJointP < 0.05).sum())
    add_body(
        doc,
        f"At the individual-session level, {n_complete} fits were complete. "
        f"Only {adjusted_any_raw} had an unadjusted adjusted-vertical 'any MS' "
        f"p < .05 and {adjusted_int_raw} had an unadjusted vertical interaction "
        f"p < .05. The 2-D vector model produced {vector_any_raw} raw any-effect "
        f"and {vector_int_raw} raw interaction findings. None survived "
        "Holm-Bonferroni correction across sessions.",
    )
    add_figure(
        doc,
        ANALYSIS_ROOT
        / "grouped_monkey_roi_ms_choice"
        / "probability_plots"
        / "grouped_ms_y_effects_95ci.png",
        4,
        "Grouped session-stratified vertical-MS results. The upper-left panel "
        "tests whether MS_y predicts choice on average; the upper-right panel "
        "tests EM moderation. The lower panels show condition-specific slopes. "
        "All adjusted 95% confidence intervals include zero.",
        "Four coefficient plots showing average vertical microsaccade slopes, "
        "MS-by-stimulation interactions, and condition-specific slopes for four "
        "monkey-area groups.",
        width=6.35,
        page_break_before=True,
    )

    add_heading(doc, "4.5 Does EM modulate the MS_y-to-choice relationship?", 2)
    add_body(
        doc,
        "The preferred adjusted interaction estimates were -0.043 in Jim-MT, "
        "-0.012 in Jim-FST, -0.045 in Clay-MT, and -0.011 in Clay-FST. Raw "
        "interaction p values ranged from 0.245 to 0.748; Holm-adjusted values "
        "ranged from 0.980 to 1.000. The adjusted 2-D X/Y interaction tests had "
        "raw p values from 0.454 to 0.813 and Holm p = 1.000 in every group. "
        "There is therefore no evidence that EM makes vertical MS more or less "
        "predictive of the later choice.",
    )

    add_heading(doc, "4.6 Pooled-model sensitivity", 2)
    pooled = pooled.set_index("Group").loc[group_order].reset_index()
    pooled_rows = []
    for _, row in pooled.iterrows():
        pooled_rows.append(
            [
                row.Group,
                f"{row.AverageSlope:.3f}",
                fmt_p(row.AverageSlopeP_Holm),
                f"{row.Interaction:.3f}",
                fmt_p(row.InteractionP_Holm),
                fmt_p(row.AnyMSEffectJointP_Holm),
            ]
        )
    add_table(
        doc,
        [
            "Group",
            "Average slope",
            "Average Holm p",
            "MS_y x EM",
            "Interaction Holm p",
            "Any MS Holm p",
        ],
        pooled_rows,
        [1300, 1500, 1650, 1350, 1760, 1800],
        alignments=[
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    add_body(
        doc,
        "The exact pooled logistic model reached the same qualitative "
        "conclusion. Average-slope Holm p values were 0.334 to 0.428; "
        "interaction Holm p values were 0.762 to 1.000; and joint any-MS Holm "
        "p values were 0.697 to 0.840. Session-stratified and pooled estimates "
        "were similar and all corresponding 95% confidence intervals crossed "
        "zero.",
    )
    add_figure(
        doc,
        ANALYSIS_ROOT
        / "grouped_monkey_roi_ms_choice"
        / "pooled_logistic"
        / "pooled_vs_session_stratified_ms_y_effects_95ci.png",
        5,
        "Sensitivity comparison of the unstratified pooled logistic model and "
        "the model with session-specific condition intercepts. Neither model "
        "supports a reliable average MS_y slope or MS_y-by-EM interaction.",
        "Two coefficient plots comparing pooled and session-stratified vertical "
        "microsaccade effects and interactions across monkey-area groups.",
        width=6.35,
        page_break_before=True,
    )

    # Integrated interpretation.
    add_heading(doc, "5. Integrated interpretation", 1)
    add_heading(doc, "5.1 What the positive session-level results mean", 2)
    add_body(
        doc,
        "The centered analysis is the strongest evidence in favor of an "
        "EM-MS-behavior link. In many sessions, an EM-induced upward shift in "
        "mean MS_y accompanied an upward shift in fitted choice bias, and a "
        "downward shift accompanied a downward bias shift. This relationship "
        "was especially coherent in MT, where both sign concordance and "
        "change-magnitude scaling were significant.",
    )
    add_body(
        doc,
        "However, this is an aggregate change-change association. It does not "
        "identify which trials produced the behavioral effect, and it cannot "
        "distinguish mediation from a shared upstream EM effect. EM could alter "
        "a latent oculomotor/decision state that independently changes mean "
        "MS_y and psychometric bias. Under that account, MS_y is a useful marker "
        "of the state without being the mechanism that generates the choice.",
    )
    add_heading(doc, "5.2 Why the trial-level null result matters", 2)
    add_body(
        doc,
        "If vertical MS were a proximal driver of the up/down decision, trials "
        "with more upward MS displacement should have a reliably different "
        "probability of Choice = 1 than trials with more downward displacement, "
        "after accounting for signed motion evidence, visual condition, EM, and "
        "recording session. That prediction was not supported. The adjusted "
        "slopes were small, their confidence intervals included zero, and the "
        "result held across pooled, session-stratified, vector, and per-session "
        "models.",
    )
    add_heading(doc, "5.3 Answer to the central question", 2)
    add_callout(
        doc,
        "Interpretation",
        "The analyses do not demonstrate that EM made vertical-MS changes that "
        "account for the monkeys' choices. They demonstrate a reproducible "
        "session-level covariation between EM-induced MS_y change and behavioral "
        "bias change, especially in MT, but not the within-trial choice-predictive "
        "link needed to treat MS_y as a behavioral mediator.",
    )
    add_body(
        doc,
        "Accordingly, the most defensible language is: 'EM-induced vertical-MS "
        "changes track EM-induced behavioral-bias changes across sessions, but "
        "vertical MS direction does not robustly predict individual choices and "
        "its choice relationship is not detectably modulated by EM. Vertical MS "
        "is therefore better interpreted as a correlated readout than as an "
        "established causal intermediary.'",
    )

    # Limitations.
    add_heading(doc, "6. Limitations and inferential boundaries", 1)
    limitation_rows = [
        [
            "MS-containing trials only",
            "Direction models exclude trials without an MS. They cannot capture a pathway through EM-induced changes in MS occurrence.",
        ],
        [
            "Aggregate bias outcomes",
            "Psychometric bias and session-mean MS_y can covary because of shared session-level factors; this is an ecological association.",
        ],
        [
            "No formal indirect effect",
            "The pipeline does not estimate a multilevel natural indirect effect or product-of-coefficients EM -> MS_y -> choice pathway.",
        ],
        [
            "Response-code sign",
            "The model predicts Choice = 1. The up/down interpretation of coefficient sign requires confirming the behavioral response coding.",
        ],
        [
            "Multiple events per trial",
            "Averaging event vectors is reproducible but may hide timing, order, or opposing MS events within a trial.",
        ],
        [
            "Temporal and latent-state confounding",
            "MS precedes the choice saccade, but both MS and choice may reflect the same EM-driven sensory or oculomotor state.",
        ],
        [
            "Session dependence",
            "The unstratified pooled model treats trials as independent. The session-stratified model is preferred, but it uses many fixed intercept terms rather than a hierarchical random-effects model.",
        ],
    ]
    add_table(
        doc,
        ["Issue", "Implication"],
        limitation_rows,
        [2200, 7160],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=8.4,
    )

    # Next analysis.
    add_heading(doc, "7. Recommended confirmatory analysis", 1)
    add_body(
        doc,
        "A confirmatory mediation analysis should preserve the multilevel "
        "structure and include every valid trial, not only trials containing an "
        "MS. Because MS direction is undefined when no MS occurs, a two-part "
        "mediator is preferable.",
    )
    next_rows = [
        [
            "Occurrence model",
            "AnyMS ~ EM + SignedCoherence * VisualCondition + (1 + EM | Session)",
            "Tests whether EM changes the probability that a trial contains an MS.",
        ],
        [
            "Direction model",
            "MS_y | AnyMS = 1 ~ EM + SignedCoherence * VisualCondition + (1 + EM | Session)",
            "Tests the EM effect on vertical direction conditional on an MS.",
        ],
        [
            "Choice model",
            "Choice ~ EM + AnyMS + MS_y + EM:MS_y + SignedCoherence * VisualCondition + session effects",
            "Estimates the within-trial occurrence and direction pathways.",
        ],
        [
            "Indirect effects",
            "Cluster bootstrap or Bayesian posterior over sessions",
            "Quantifies uncertainty in EM -> AnyMS -> choice and EM -> MS_y -> choice pathways.",
        ],
    ]
    add_table(
        doc,
        ["Component", "Model sketch", "Purpose"],
        next_rows,
        [1650, 4800, 2910],
        alignments=[None, None, None],
        font_size=8.2,
    )
    add_body(
        doc,
        "The confirmatory analysis should predefine the response-code mapping, "
        "the MS time window relative to sensory evidence and decision formation, "
        "and whether MS_y is summarized by the first event, the last event, a "
        "time-weighted mean, or the current within-trial vector average. Session-"
        "clustered resampling should be used so trial count does not overstate "
        "precision.",
    )

    # Reproducibility.
    add_heading(doc, "8. Reproducibility", 1)
    add_body(
        doc,
        "The report was generated from the completed 12 ms, no-smoothing "
        "population output. Re-running should start from the repository root and "
        "use the MATLAB entry points below in order.",
    )
    repro_rows = [
        [
            "Population postprocessing",
            "run_12ms_population_postprocessing.m",
            "MS metric histograms, paired bias LME, centered DeltaMS_y-DeltaBias tests",
        ],
        [
            "Per-session choice models",
            "run_12ms_trialwise_ms_choice.m",
            "Trial table, 254 session summaries, coefficients, diagnostic figures",
        ],
        [
            "Grouped and pooled models",
            "run_grouped_monkey_roi_ms_choice.m",
            "Four monkey-area fits, probability plots, pooled sensitivity comparison",
        ],
    ]
    add_table(
        doc,
        ["Stage", "Entrypoint", "Primary outputs"],
        repro_rows,
        [2100, 3000, 4260],
        alignments=[None, None, None],
        font_size=8.2,
    )
    add_heading(doc, "8.1 Core analysis functions", 2)
    for line in (
        "analyze_microsaccades.m",
        "analyze_ms_behavior_bias.m",
        "analyze_paired_ms_bias_lme.m",
        "analyze_centered_ms_bias_changes.m",
        "analyze_trialwise_ms_choice.m",
        "analyze_grouped_ms_choice.m",
        "analyze_pooled_logistic_ms_choice.m",
    ):
        add_code_line(doc, line)
    add_heading(doc, "8.2 Key machine-readable result files", 2)
    output_rows = [
        [
            "EM effects on MS",
            str(
                POP_ROOT
                / "population_microsaccade_stim_nonstim_tests_complete.csv"
            ),
        ],
        [
            "Session-condition MS metrics",
            str(
                ANALYSIS_ROOT
                / "histograms"
                / "population_ms_session_condition_metrics.csv"
            ),
        ],
        [
            "Paired bias LME",
            str(
                ANALYSIS_ROOT
                / "paired_ms_y_bias_lme"
                / "paired_ms_y_bias_lme_summary.csv"
            ),
        ],
        [
            "Centered changes",
            str(
                ANALYSIS_ROOT
                / "paired_ms_y_bias_lme"
                / "centered_change"
                / "centered_delta_ms_y_bias_test_summary.csv"
            ),
        ],
        [
            "Session choice models",
            str(
                ANALYSIS_ROOT
                / "trialwise_ms_choice"
                / "trialwise_ms_choice_session_summary.csv"
            ),
        ],
        [
            "Grouped choice models",
            str(
                ANALYSIS_ROOT
                / "grouped_monkey_roi_ms_choice"
                / "grouped_ms_choice_summary.csv"
            ),
        ],
        [
            "Pooled sensitivity",
            str(
                ANALYSIS_ROOT
                / "grouped_monkey_roi_ms_choice"
                / "pooled_logistic"
                / "pooled_logistic_ms_choice_summary.csv"
            ),
        ],
    ]
    add_table(
        doc,
        ["Result", "File"],
        output_rows,
        [2050, 7310],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=7.5,
    )

    # Appendix with all cue-level aggregate results.
    doc.add_page_break()
    add_heading(doc, "Appendix A. Cue-specific aggregate results", 1)
    centered_rows = []
    for _, row in centered.sort_values(["ROI", "CueIndex"]).iterrows():
        centered_rows.append(
            [
                row.ROI,
                row.CueName,
                f"{int(row.NSessions)}",
                f"{100 * row.ConcordanceRate:.1f}%",
                fmt_ci(row.ConcordanceCI95Lower, row.ConcordanceCI95Upper),
                fmt_p(row.QuadrantP_Q_FDR_All),
                f"{row.OriginSlope:.3f}",
                fmt_p(row.OriginSlopeP_Q_FDR_All),
            ]
        )
    add_table(
        doc,
        [
            "ROI",
            "Cue",
            "N",
            "Concordance",
            "95% CI",
            "Quadrant q",
            "Origin slope",
            "Slope q",
        ],
        centered_rows,
        [650, 1150, 650, 1300, 1550, 1250, 1400, 1410],
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        font_size=7.7,
    )
    add_source_note(
        doc,
        "Quadrant q and slope q use Benjamini-Hochberg FDR correction across all "
        "eight ROI-by-cue tests.",
    )

    lme_all_rows = []
    for _, row in paired.sort_values(["ROI", "CueIndex"]).iterrows():
        lme_all_rows.append(
            [
                row.ROI,
                row.CueName,
                f"{int(row.NSessions)}",
                f"{row.CommonSlope:.3f} +/- {row.CommonSlopeSE:.3f}",
                fmt_p(row.CommonSlopeQ_FDR),
                f"{row.InteractionSlopeDifference:.3f}",
                fmt_p(row.InteractionQ_FDR),
            ]
        )
    add_table(
        doc,
        [
            "ROI",
            "Cue",
            "N",
            "Common slope +/- SE",
            "Slope q",
            "Interaction",
            "Interaction q",
        ],
        lme_all_rows,
        [700, 1250, 650, 2200, 1250, 1500, 1810],
        alignments=[
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
        font_size=7.8,
    )
    add_source_note(
        doc,
        "LME q values are FDR-corrected across the four cue definitions within "
        "each ROI and test family.",
    )

    # Document properties.
    props = doc.core_properties
    props.title = "Can EM-Driven Vertical Microsaccades Account for Up/Down Choices?"
    props.subject = "Microsaccade mediation and EM modulation analysis"
    props.author = "Microsaccade analysis repository"
    props.keywords = "microsaccades, electrical microstimulation, choice, mediation, MT, FST"
    props.comments = (
        "Generated from the completed 12 ms no-smoothing population analysis."
    )

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_report()
    print(output)
