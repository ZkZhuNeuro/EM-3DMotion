import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(r"C:\EM\PatternTuning\PatternTuning_Analysis_Summary.docx")
RESULT_ROOT = Path(r"C:\EM\PatternTuning\Final_128")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
BODY = "202124"
MUTED = "5F6368"
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "D9DEE5"
TABLE_TEXT_SIZE = 9.5  # Named compact numeric-table override.


def set_run_font(run, size=None, color=BODY, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, size, color, bold=False):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = cell_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[col_idx])


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    p.paragraph_format.keep_with_next = True
    return p


def add_labeled_paragraph(doc, label, text, *, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"{label} ")
    set_run_font(r, 11, DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, 11, BODY)
    return p


def add_callout(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.line_spacing = 1.1
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "14")
    left.set(qn("w:color"), BLUE)
    left.set(qn("w:space"), "6")
    borders.append(left)
    p_pr.append(borders)
    r = p.add_run(f"{label} ")
    set_run_font(r, 10.5, DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, 10.5, BODY)
    return p


def read_results():
    files = {
        "Lo": RESULT_ROOT / "Lo" / "LoFSTPDIAndBODI128ZeroTests.csv",
        "EM": RESULT_ROOT / "EM" / "EMFSTPDIAndBODI128ZeroTests.csv",
        "Combined": RESULT_ROOT / "Combined" / "CombinedLoEMFSTPDIAndBODI128ZeroTests.csv",
    }
    records = []
    for dataset, path in files.items():
        with path.open(newline="", encoding="utf-8-sig") as handle:
            for row in csv.DictReader(handle):
                records.append(
                    {
                        "Dataset": dataset,
                        "Metric": row["Metric"],
                        "N": int(float(row["N"])),
                        "Mean": float(row["Mean"]),
                        "Median": float(row["Median"]),
                        "PValue": float(row["PValue"]),
                        "PValue_Bonferroni": float(row["PValue_Bonferroni"]),
                    }
                )
    return records


def format_p(value):
    return f"{value:.2e}" if value < 0.001 else f"{value:.3f}"


def add_results_table(doc, records):
    headers = ("Dataset", "Metric", "n", "Mean", "Median", "Rank-sum p", "Bonferroni p")
    widths = [1600, 1100, 600, 1100, 1100, 1800, 2060]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for col_idx, text in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        set_cell_shading(cell, HEADER_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, TABLE_TEXT_SIZE, DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])

    for record in records:
        values = (
            record["Dataset"],
            record["Metric"],
            str(record["N"]),
            f'{record["Mean"]:.3f}',
            f'{record["Median"]:.3f}',
            format_p(record["PValue"]),
            format_p(record["PValue_Bonferroni"]),
        )
        cells = table.add_row().cells
        for col_idx, text in enumerate(values):
            cells[col_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, TABLE_TEXT_SIZE, BODY)
    set_table_geometry(table, widths)


def build_document():
    records = read_results()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 11, BODY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    h1 = styles["Heading 1"]
    set_style_font(h1, 16, BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2, 13, BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("PATTERN TUNING | FINAL 1.28 ANALYSIS")
    set_run_font(r, 8.5, MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("12 August 2026")
    set_run_font(r, 8.5, MUTED)

    # memo_masthead with a compact title override for a one-page technical brief.
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(3)
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("FST Preference and Binocular Optic Flow Discrimination")
    set_run_font(r, 21, "000000", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.keep_with_next = True
    r = subtitle.add_run("Concise summary | Lo, EM, and combined FST datasets")
    set_run_font(r, 12.5, MUTED)

    add_callout(
        doc,
        "Metrics.",
        "PDI uses TDI for toward-preferring neurons and ADI for away-preferring neurons, "
        "evaluated at the matched slow 2D speed (about 4.2 deg/s). BODI = "
        "(Combined FR - Stereo FR) / (Combined FR + Stereo FR), comparing cue 1 with "
        "cue 4 at coherence +1 for toward preference and -1 for away preference.",
    )

    add_heading(doc, "Selection and inference")
    add_labeled_paragraph(
        doc,
        "Lo:",
        'ROI == "FST", significant CLR ANOVA, and Z_quad == 2 (n = 58).',
    )
    add_labeled_paragraph(
        doc,
        "EM:",
        'ROI == "FST", ND == "3D", and Z3D_v_Z2D > 1.28 (n = 24). The combined sample contains 82 neurons.',
    )
    add_labeled_paragraph(
        doc,
        "Test:",
        "two-sided Wilcoxon rank-sum against an equal-sized zero reference. "
        "Bonferroni correction was applied across PDI and BODI within each dataset.",
        after=7,
    )

    add_heading(doc, "Results")
    add_results_table(doc, records)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(7)
    r = note.add_run("Note. ")
    set_run_font(r, 8.5, MUTED, bold=True)
    r = note.add_run("Histograms pool toward- and away-preferring neurons; preference is used only to choose the PDI component and BODI coherence endpoint.")
    set_run_font(r, 8.5, MUTED)

    add_heading(doc, "Summary")
    add_labeled_paragraph(
        doc,
        "Population shift:",
        "PDI and BODI means and medians were positive for Lo, EM, and the combined sample.",
    )
    add_labeled_paragraph(
        doc,
        "Statistical result:",
        "all six pooled distributions differed from zero after the two-metric Bonferroni correction.",
        after=0,
    )

    doc.core_properties.title = "FST PDI and BODI - Final 1.28 Analysis Summary"
    doc.core_properties.subject = "Concise summary of strict-criterion Lo and EM FST analyses"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "PDI, BODI, optic flow, FST, Lo, EM, 1.28"
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
